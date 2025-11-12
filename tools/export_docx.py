#!/usr/bin/env python3
import sys, pathlib
from docx import Document
from docx.shared import Pt

def main():
    src = pathlib.Path(sys.argv[1] if len(sys.argv)>1 else "outputs/final_bullets.txt")
    dst = pathlib.Path(sys.argv[2] if len(sys.argv)>2 else "outputs/final_bullets.docx")
    if not src.exists():
        print(f"ERROR: {src} not found"); sys.exit(2)
    lines = [l.strip() for l in src.read_text(encoding="utf-8").splitlines() if l.strip()]
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    for L in lines:
        doc.add_paragraph().add_run("• " + L)
    doc.save(dst.as_posix())
    print(f"Wrote {dst}")
if __name__ == "__main__":
    main()
